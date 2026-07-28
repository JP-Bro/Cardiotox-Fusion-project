"""
create_docx.py -- Generate a highly professional Microsoft Word Document (.docx)
for the Week 1 deliverable, containing embedded tables, images, Jupyter execution logs,
and a scientific discussion addressing comparative evaluation of GNN vs. Transformer models.
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from config import CFG

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell margins (padding) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, data, column_widths=None):
    """Add a stylized table with header highlighting and light gray gridlines."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False

    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], "1A1A2E")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)

    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"
        # Highlight total row
        if "Total" in str(row_data[0]):
            bg_color = "EAEAEA"
            
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = str(text)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                if "Total" in str(row_data[0]):
                    run.font.bold = True

    # Adjust widths
    if column_widths:
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = Inches(width)
                
    # Add a blank spacing paragraph after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    
    # Configure document margins (1 inch / 2.54 cm all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style configurations
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11.5)
    font.color.rgb = RGBColor(0, 0, 0)

    # Title Block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Cardiotox-Fusion: Comparative Evaluation of GNN and Transformer Architectures for Predicting FDA DICTrank Cardiotoxicity")
    title_run.font.bold = True
    title_run.font.size = Pt(15)
    title_run.font.name = "Times New Roman"
    title.paragraph_format.space_after = Pt(6)

    # Metadata Info
    meta = doc.add_paragraph()
    meta.paragraph_format.line_spacing = 1.3
    meta.paragraph_format.space_after = Pt(18)
    
    runs_data = [
        ("Milestone: ", True), ("Week 1 Deliverable   |   ", False),
        ("Date: ", True), ("July 28, 2026\n", False),
        ("Team Members: ", True), ("Jimit Patel, Vineela, Aditya, Nehal, Pranjal\n", False),
        ("Submitted to: ", True), ("Dr. Bharat Manna (School of Bio Sciences and Technology, VIT)\n", False),
        ("Repository: ", True), ("github.com/JP-Bro/Cardiotox-Fusion-project", False)
    ]
    for text, is_bold in runs_data:
        r = meta.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.bold = is_bold
        r.font.color.rgb = RGBColor(50, 50, 50)

    # 1. Dataset Usability Audit
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Dataset Usability Audit")
    r.font.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 0, 0)

    p1 = doc.add_paragraph(
        "To establish the modeling boundaries for our comparative evaluation, "
        "we performed a joint audit of the FDA DICTrank labeling set and the LINCS L1000 GSE70138 "
        "level-5 expression matrix. Compounds were retained in the usable cohort only if they met "
        "the following criteria:"
    )
    p1.paragraph_format.space_after = Pt(6)

    bullets = [
        "Possessed a valid, parseable chemical structure (verified via RDKit).",
        "Matched a unique Level-5 perturbation signature (COMPZ) in the HA1E cell line (10.0 µM dose, 24-hour time point).",
        "Carried a non-ambiguous DICTrank label mapping to binary classification (No concern = 0; Less/Most concern = 1)."
    ]
    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        br = bp.add_run(b)
        br.font.size = Pt(11)

    p2 = doc.add_paragraph(
        "Out of the raw 1,211 labeled DICTrank compounds, a total of 562 compounds met all three requirements, "
        "forming our core dataset. The breakdown across the cardiotoxicity label categories is detailed below:"
    )
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(10)

    # Table 1: Usability Audit Counts
    t1_headers = ["DICTrank Category", "Model Label", "Usable Count", "Percentage"]
    t1_data = [
        ["Most Concern (High Risk)", "1 (Positive)", "199", "35.4%"],
        ["Less Concern (Medium Risk)", "1 (Positive)", "246", "43.8%"],
        ["No Concern (Safe)", "0 (Negative)", "117", "20.8%"],
        ["Total Usable Set", "—", "562", "100.0%"]
    ]
    add_styled_table(doc, t1_headers, t1_data, column_widths=[2.8, 1.2, 1.2, 1.3])

    # Jupyter Execution Log box (Audit section)
    p_log_title = doc.add_paragraph()
    r_lt = p_log_title.add_run("Jupyter Notebook Audit Execution Verification Log:")
    r_lt.font.bold = True
    r_lt.font.size = Pt(10)
    p_log_title.paragraph_format.space_before = Pt(6)
    p_log_title.paragraph_format.space_after = Pt(3)

    log_box1 = doc.add_paragraph()
    set_cell_background(doc.add_table(rows=1, cols=1).rows[0].cells[0], "F5F5F7")
    table_log1 = doc.tables[-1]
    table_log1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_log1.rows[0].cells[0].width = Inches(5.8)
    set_cell_margins(table_log1.rows[0].cells[0], top=80, bottom=80, left=120, right=120)
    p_code1 = table_log1.rows[0].cells[0].paragraphs[0]
    p_code1.paragraph_format.line_spacing = 1.15
    p_code1.paragraph_format.space_after = Pt(0)
    r_code1 = p_code1.add_run(
        "In [2]: # Load and merge datasets\n"
        "Raw DICTrank count: 1211\n"
        "L1000 matched compounds: 1048\n"
        "Merged set count: 1048\n\n"
        "In [4]: # Class Counts Verification\n"
        "--- CLASS COUNTS (DICTrank Full Set) ---\n"
        "  less           : 451\n"
        "  most           : 318\n"
        "  no             : 279\n\n"
        "--- CLASS COUNTS (LINCS-Overlapping Usable Set) ---\n"
        "  less           : 246 (43.8%)\n"
        "  most           : 199 (35.4%)\n"
        "  no             : 117 (20.8%)"
    )
    r_code1.font.name = "Courier New"
    r_code1.font.size = Pt(8.5)
    r_code1.font.color.rgb = RGBColor(30, 30, 30)

    # Insert spacing paragraph
    p_space_after_log = doc.add_paragraph()
    p_space_after_log.paragraph_format.space_after = Pt(12)

    # Insert Image 1: Usable Cohort Plot
    img1_path = os.path.join(CFG.RESULTS_DIR, "phase1_usable_cohort.png")
    if os.path.isfile(img1_path):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.paragraph_format.space_after = Pt(18)
        p_img1.add_run().add_picture(img1_path, width=Inches(4.5))
        caption1 = doc.add_paragraph()
        caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption1.paragraph_format.space_after = Pt(12)
        cr1 = caption1.add_run("Figure 1: Usable Dataset Cohort distribution by FDA DICTrank concern class.")
        cr1.font.italic = True
        cr1.font.size = Pt(9.5)

    # 2. Leakage-Free Dataset Splits
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    r2 = h2.add_run("2. Leakage-Free Dataset Splits")
    r2.font.bold = True
    r2.font.size = Pt(13)

    p3 = doc.add_paragraph(
        "To prevent target leakage and ensure realistic validation generalizability, "
        "we implemented two separate stratified splitting schemes. In both schemes, splits are frozen at a 70 / 15 / 15 ratio. "
        "For scaffold splits, whole carbon scaffold clusters are assigned to the held-out test set until at least 15% of unique compounds are covered."
    )
    p3.paragraph_format.space_after = Pt(12)

    # Heading 2.1: Drug-Level Splits
    h2_1 = doc.add_paragraph()
    h2_1.paragraph_format.space_before = Pt(8)
    h2_1.paragraph_format.space_after = Pt(6)
    r2_1 = h2_1.add_run("A. Drug-Level Splits (Stratified)")
    r2_1.font.bold = True
    r2_1.font.italic = True
    r2_1.font.size = Pt(11.5)

    # Table 2: Drug Splits
    t2_headers = ["Partition", "Compound Count", "Positive (Tox) Count", "Negative (Safe) Count", "Positive Rate"]
    t2_data = [
        ["Train", "393", "311", "82", "79.1%"],
        ["Validation", "84", "67", "17", "79.8%"],
        ["Test", "85", "67", "18", "78.8%"],
        ["Total Set", "562", "445", "117", "79.2%"]
    ]
    add_styled_table(doc, t2_headers, t2_data, column_widths=[1.5, 1.3, 1.3, 1.3, 1.1])

    # Heading 2.2: Scaffold-Level Splits
    h2_2 = doc.add_paragraph()
    h2_2.paragraph_format.space_before = Pt(8)
    h2_2.paragraph_format.space_after = Pt(6)
    r2_2 = h2_2.add_run("B. Bemis-Murcko Scaffold Splits")
    r2_2.font.bold = True
    r2_2.font.italic = True
    r2_2.font.size = Pt(11.5)

    # Table 3: Scaffold Splits
    t3_data = [
        ["Train", "393", "312", "81", "79.4%"],
        ["Validation", "84", "63", "21", "75.0%"],
        ["Test", "85", "70", "15", "82.4%"],
        ["Total Set", "562", "445", "117", "79.2%"]
    ]
    add_styled_table(doc, t2_headers, t3_data, column_widths=[1.5, 1.3, 1.3, 1.3, 1.1])

    # Jupyter Execution Log box (Splits section)
    p_log_title2 = doc.add_paragraph()
    r_lt2 = p_log_title2.add_run("Jupyter Notebook Splits Partition Verification Log:")
    r_lt2.font.bold = True
    r_lt2.font.size = Pt(10)
    p_log_title2.paragraph_format.space_before = Pt(6)
    p_log_title2.paragraph_format.space_after = Pt(3)

    table_log2 = doc.add_table(rows=1, cols=1)
    table_log2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_log2.rows[0].cells[0].width = Inches(5.8)
    set_cell_background(table_log2.rows[0].cells[0], "F5F5F7")
    set_cell_margins(table_log2.rows[0].cells[0], top=80, bottom=80, left=120, right=120)
    p_code2 = table_log2.rows[0].cells[0].paragraphs[0]
    p_code2.paragraph_format.line_spacing = 1.15
    p_code2.paragraph_format.space_after = Pt(0)
    r_code2 = p_code2.add_run(
        "In [5]: # Generate and verify splits\n"
        "Drug-level partitions:\n"
        "  train   : 393 compounds (pos rate: 79.1%)\n"
        "  test    : 85 compounds (pos rate: 78.8%)\n"
        "  val     : 84 compounds (pos rate: 79.8%)\n\n"
        "Scaffold-level partitions:\n"
        "  train   : 393 compounds (pos rate: 79.4%)\n"
        "  test    : 85 compounds (pos rate: 82.4%)\n"
        "  val     : 84 compounds (pos rate: 75.0%)"
    )
    r_code2.font.name = "Courier New"
    r_code2.font.size = Pt(8.5)
    r_code2.font.color.rgb = RGBColor(30, 30, 30)

    # Insert spacing paragraph
    p_space_after_log2 = doc.add_paragraph()
    p_space_after_log2.paragraph_format.space_after = Pt(12)

    # Insert Image 2: Splits Distribution Plot
    img2_path = os.path.join(CFG.RESULTS_DIR, "phase1_splits_distribution.png")
    if os.path.isfile(img2_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.space_after = Pt(18)
        p_img2.add_run().add_picture(img2_path, width=Inches(4.5))
        caption2 = doc.add_paragraph()
        caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption2.paragraph_format.space_after = Pt(12)
        cr2 = caption2.add_run("Figure 2: Target partition sizes and positive rates across drug-level and scaffold splits.")
        cr2.font.italic = True
        cr2.font.size = Pt(9.5)

    # 3. Methodological Details & Scientific Constraints
    h3_meth = doc.add_paragraph()
    h3_meth.paragraph_format.space_before = Pt(18)
    h3_meth.paragraph_format.space_after = Pt(6)
    r3_meth = h3_meth.add_run("3. Methodology, Signature Aggregation, and Dual Output System")
    r3_meth.font.bold = True
    r3_meth.font.size = Pt(13)

    p_meth1 = doc.add_paragraph(
        "To ensure mathematical rigor and prevent data leakage or validation contamination, "
        "we have defined explicit constraints for the dataset preprocessing and dual-output inference pipeline:"
    )
    p_meth1.paragraph_format.space_after = Pt(6)

    meth_points = [
        ("L1000 Signature Aggregation Rule: ", True,
         "A single drug may have multiple replicate signatures measured across different cell lines, doses, and time points in LINCS L1000. "
         "To handle this, we restrict our selection to Level-5 signatures (COMPZ z-scores) in the HA1E cell line. "
         "For any drug with multiple replicate signatures matching 10.0 µM dose and 24 h exposure, "
         "we mean-aggregate the z-score vectors into a single, consolidated 978-dimensional transcriptomic representation per drug."),
         
        ("Dual Independent Outputs (GNN & Transformer): ", True,
         "We evaluate and display two independent output scores for every compound: "
         "(1) Structural Risk Score (from GNN), and (2) Biological Perturbation Score (from Transformer). "
         "No black-box fusion is performed, ensuring transparent, un-contaminated evaluation across both modalities."),

        ("Mathematical Imputation Tagging & Disclaimer: ", True,
         "For novel or out-of-dataset compounds lacking experimental L1000 signatures, an MLP imputer calculates a mathematical estimation "
         "of the gene expression vector from chemical structure. Crucially, the system explicitly tags the output: "
         "'[Notice] Imputed Signature: This compound uses a mathematically calculated gene profile. Mathematically estimated results may vary from experimental observations.' "
         "Furthermore, imputed profiles are reported separately and excluded from the core experimental benchmark.")
    ]
    for title, is_bold, text in meth_points:
        dp = doc.add_paragraph(style='List Bullet')
        dp.paragraph_format.space_after = Pt(4)
        run_title = dp.add_run(title)
        run_title.font.bold = is_bold
        run_title.font.size = Pt(11)
        run_text = dp.add_run(text)
        run_text.font.size = Pt(11)

    # 4. Scientific Discussion & Publishable Novelty
    h4_disc = doc.add_paragraph()
    h4_disc.paragraph_format.space_before = Pt(18)
    h4_disc.paragraph_format.space_after = Pt(6)
    r4_disc = h4_disc.add_run("4. Project Discussion: Model Comparison, Leakage, and Academic Novelty")
    r4_disc.font.bold = True
    r4_disc.font.size = Pt(13)

    p_disc1 = doc.add_paragraph(
        "A central question of this research is comparing the predictive performance of a "
        "purely structure-based model (GNN) against a purely transcriptomics-based model (Transformer) "
        "on clinical cardiotoxicity labels. Based on the dataset properties and structural alerts, "
        "we present the following analysis and academic novelty strategy:"
    )
    p_disc1.paragraph_format.space_after = Pt(6)

    disc_points = [
        ("DICTrank vs. hERG Channel Blockade: ", True,
         "We clarify that DICTrank represents general cardiotoxicity concern classes, not just hERG channel blockade. "
         "While hERG blockade is the primary physical mechanism of cardiotoxicity, DICTrank encompasses other forms of toxicities. "
         "Therefore, our model is trained to predict general cardiotoxicity (DICTrank classes), and hERG blockade is discussed as the principal "
         "mechanistic subclass, rather than being claimed as a direct prediction target."),
         
        ("Bemis-Murcko Scaffold Splits as Primary Novelty: ", True,
         "Prior publications (e.g. Seal et al.) evaluated models using standard random or stratified splits. "
         "This project introduces Bemis-Murcko scaffold splits on DICTrank, assigning whole scaffold clusters to the test set until at least 15% of unique drugs are held out. "
         "This evaluation directly quantifies how much cardiotoxicity prediction is due to chemical-series memorization versus out-of-distribution generalizability. "
         "This scaffold-split analysis constitutes our primary publishable novelty."),
         
        ("Testing the 'Representation Extraction' Hypothesis: ", True,
         "The near-random performance of L1000 signatures in literature (~0.57 AUC) raises the question of whether transcriptomic profiles lack signal "
         "or if previous shallow classifiers failed to extract it. By utilizing a multi-head Transformer encoder on the experimental z-score sequence, "
         "we directly test whether a learned representation can recover predictive biological signal that simpler methods missed. "
         "Evaluating this cleanly, with transparent imputation tagging, provides a high-quality benchmarking study.")
    ]
    
    for title, is_bold, text in disc_points:
        dp = doc.add_paragraph(style='List Bullet')
        dp.paragraph_format.space_after = Pt(4)
        run_title = dp.add_run(title)
        run_title.font.bold = is_bold
        run_title.font.size = Pt(11)
        run_text = dp.add_run(text)
        run_text.font.size = Pt(11)

    # 5. Verification & Reproducibility
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    r5 = h5.add_run("5. Verification and Reproducibility")
    r5.font.bold = True
    r5.font.size = Pt(13)

    p5 = doc.add_paragraph(
        "All data processing steps are fully automated. The random seed is frozen at 42 in config.py, "
        "and both split files are frozen and saved in the repository at data/splits/drug_split.csv "
        "and data/splits/scaffold_split.csv. "
        "These splits will remain frozen throughout the modeling pipeline."
    )
    p5.paragraph_format.space_after = Pt(12)

    # Save document with locked file fallback
    out_docx_path = os.path.join(CFG.RESULTS_DIR, "Cardiotox_Fusion_Phase1_Report.docx")
    try:
        doc.save(out_docx_path)
        print(f"Word document saved successfully to: {out_docx_path}")
    except PermissionError:
        fallback_path = os.path.join(CFG.RESULTS_DIR, "Cardiotox_Fusion_Phase1_Report_v2.docx")
        doc.save(fallback_path)
        print(f"WARNING: Primary file locked. Word document saved to fallback: {fallback_path}")
        print("Please close Microsoft Word to allow overwriting of the primary file.")

if __name__ == "__main__":
    main()
